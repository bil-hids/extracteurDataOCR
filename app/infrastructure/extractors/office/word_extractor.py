"""Extracteur Word utilisant python-docx."""

import asyncio
from typing import Any

from docx import Document

from app.core.exceptions import ExtractionFailedError
from app.core.logging import get_logger
from app.domain.value_objects.content_metadata import ContentMetadata
from app.domain.value_objects.extraction_result import ExtractionResult, TableBlock, TextBlock
from app.infrastructure.extractors.base import BaseExtractor


class WordExtractor(BaseExtractor):
    """Extracteur Word avec python-docx."""

    def __init__(self, logger: Any = None) -> None:
        """Initialiser l'extracteur."""
        super().__init__(logger or get_logger(__name__))

    async def extract(self, file_path: str) -> ExtractionResult:
        """Extraire le contenu d'un fichier Word."""
        self._validate_file(file_path)

        try:
            loop = asyncio.get_event_loop()
            result = await loop.run_in_executor(None, self._extract_sync, file_path)
            return result
        except Exception as e:
            self.logger.exception(f"Erreur lors de l'extraction Word: {e}")
            raise ExtractionFailedError(f"Échec de l'extraction Word: {str(e)}")

    def _extract_sync(self, file_path: str) -> ExtractionResult:
        """Extraction synchrone."""
        text_blocks: list[TextBlock] = []
        tables: list[TableBlock] = []

        doc = Document(file_path)

        # Extraire métadonnées
        core_props = doc.core_properties
        metadata = {
            "title": core_props.title,
            "author": core_props.author,
            "created": str(core_props.created) if core_props.created else None,
            "modified": str(core_props.modified) if core_props.modified else None,
        }

        # Extraire le texte avec structure
        current_section = None
        order = 0

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue

            # Détecter les titres
            style_name = paragraph.style.name if paragraph.style else ""
            is_heading = "Heading" in style_name or paragraph.style.name.startswith("Heading")

            if is_heading:
                current_section = text
                level = self._extract_heading_level(style_name)

                metadata_obj = ContentMetadata(
                    order=order,
                    section_title=text,
                    section_level=level,
                    extraction_method="python-docx",
                    additional_metadata={"style": style_name},
                )
                text_blocks.append(TextBlock(content=text, metadata=metadata_obj))
                order += 1
            else:
                # Paragraphe normal
                metadata_obj = ContentMetadata(
                    order=order,
                    section_title=current_section,
                    extraction_method="python-docx",
                )
                text_blocks.append(TextBlock(content=text, metadata=metadata_obj))
                order += 1

        # Extraire les tableaux
        for table in doc.tables:
            headers: list[str] = []
            rows: list[list[str]] = []

            for i, row in enumerate(table.rows):
                # Nettoyer les cellules : convertir None en chaîne vide
                row_data = [
                    cell.text.strip() if cell.text is not None else ""
                    for cell in row.cells
                ]
                if i == 0:
                    headers = row_data
                else:
                    rows.append(row_data)

            if headers or rows:
                metadata_obj = ContentMetadata(
                    order=len(tables),
                    extraction_method="python-docx",
                )
                tables.append(TableBlock(headers=headers, rows=rows, metadata=metadata_obj))

        return ExtractionResult(
            text_blocks=text_blocks, tables=tables, raw_metadata=metadata
        )

    def _extract_heading_level(self, style_name: str) -> int:
        """Extraire le niveau de titre depuis le nom de style."""
        if "Heading" in style_name:
            try:
                # Essayer d'extraire le numéro (Heading 1, Heading 2, etc.)
                level_str = style_name.split()[-1]
                return int(level_str)
            except (ValueError, IndexError):
                pass
        return 1  # Par défaut

    async def extract_tables(self, file_path: str) -> list[Any]:
        """Extraire uniquement les tableaux."""
        result = await self.extract(file_path)
        return result.tables

    async def extract_images(self, file_path: str) -> list[Any]:
        """python-docx peut extraire les images mais c'est complexe."""
        return []

    def supports(self, file_type: str) -> bool:
        """Vérifier si le type Word est supporté."""
        return file_type in [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        ]

